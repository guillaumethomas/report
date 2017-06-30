from collections import defaultdict
from tabulate import tabulate
from prettytable import PrettyTable
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
#https://python-docx.readthedocs.io/en/latest/

def essai():
    print('Essai')

    essai_lst = [{'name': 'guillaume', 'id': 150}, {'name': 'ben','id': 180},\
                 {'name': 'guillaume', 'id': 180}]


    #print(essai_lst)

    name_dict = defaultdict(list)
    nb_dict = defaultdict(int)

    for dc in essai_lst:
        name_dict[dc['name']].append(str(dc['id']))
        nb_dict[dc['name']] += 1

    name_lst = list(nb_dict.keys())

    return name_dict, nb_dict, name_lst

def table(data_lst):

    name_dict, nb_dict, name_lst = data_lst

    table = PrettyTable(['Name','Bugs'])
    table_2 = PrettyTable(['Name', 'Bugs lists'])
    table_3 = PrettyTable(['Name','# Bugs', 'Bugs lists'])

    for name in name_lst:
        table.add_row([name,nb_dict[name]])
        table_2.add_row([name, ','.join(name_dict[name])])
        table_3.add_row([name, nb_dict[name],','.join(name_dict[name])])

    print(table)
    print(table_2)
    print(table_3)


    tables_txt = [table.get_string(), table_2.get_string(),table_3.get_string()]

    file = open('report.txt', 'w')

    for txt in tables_txt:
        file.write(txt)
        file.write('\n')


    file.close()

    return file

def chart(x,y):

    plt.plot(x,y, 'ro')
    plt.axis([0, 6, 0, 20])
    return plt

def word_doc(name,pic_lst):
    document = Document()

    document.add_heading('Bug report', 0)

    p = document.add_paragraph('A plain paragraph having some ')

    document.add_picture(pic_lst[0], width=Inches(5))
    return document.save(name)

if __name__ == "__main__":
    data = essai()

    file = table(data)
    print(file.name)

    fig = chart([1, 2, 3, 4], [1, 4, 9, 16])
    pic = [fig.savefig('test.png')]
    word_doc('test.docx',['test.png'])
    fig.show()
    print(type(pic[0]))
